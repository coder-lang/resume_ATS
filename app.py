# app.py - Enhanced ATS Resume Builder with Best Practices
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import io
import json
import PyPDF2
from dotenv import load_dotenv
from openai import OpenAI

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key')

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found")

client = OpenAI(api_key=OPENAI_API_KEY)
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# ---------------------------
# File parsing
# ---------------------------
def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return '\n'.join([p.text for p in doc.paragraphs])

# ---------------------------
# AI extraction with streaming support
# ---------------------------
def extract_resume_data_with_ai(resume_text: str) -> dict:
    max_length = 6000
    if len(resume_text) > max_length:
        resume_text = resume_text[:max_length] + "\n...(truncated)"
    
    prompt = f"""Extract resume data and return ONLY valid JSON:

{{
  "personal": {{
    "name": "Full Name",
    "email": "email",
    "phone": "phone",
    "location": "City, State",
    "linkedin": "URL or null",
    "github": "URL or null",
    "portfolio": "URL or null"
  }},
  "summary": "Professional summary or null",
  "education": [
    {{
      "institution": "University",
      "degree": "Degree",
      "field": "Major",
      "location": "City, State",
      "graduation": "Year",
      "gpa": "GPA or null",
      "achievements": ["Achievement 1"]
    }}
  ],
  "experience": [
    {{
      "company": "Company Name",
      "position": "Job Title",
      "location": "City, State",
      "start": "Date",
      "end": "Date or Present",
      "description": "Brief description or null",
      "achievements": ["Key achievement 1", "Key achievement 2"]
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description": "Brief description",
      "technologies": ["Tech1", "Tech2"],
      "start": "Date or null",
      "end": "Date or null",
      "achievements": ["Achievement 1"]
    }}
  ],
  "skills": {{
    "technical": ["Python", "React"],
    "soft": ["Leadership", "Communication"],
    "languages": ["English (Native)", "Spanish (Intermediate)"],
    "tools": ["Git", "Docker"],
    "additional": ["Additional skill 1"]
  }},
  "certifications": [
    {{
      "name": "Certification Name",
      "issuer": "Issuing Organization",
      "date": "Date or null",
      "credential": "Credential ID or null"
    }}
  ],
  "leadership": [
    {{
      "organization": "Organization Name",
      "role": "Position",
      "location": "City, State or null",
      "start": "Date",
      "end": "Date or Present",
      "achievements": ["Achievement 1"]
    }}
  ]
}}

Resume: {resume_text}

Return ONLY valid JSON without trailing commas."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a resume parser. Return valid JSON. Use null for missing data."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=2000,
            timeout=25
        )
        result = response.choices[0].message.content.strip()
        
        # Clean response
        if result.startswith("```"):
            parts = result.split("```")
            result = parts[1] if len(parts) > 1 else result
        result = result.strip().lstrip("json").strip()
        
        # Remove trailing commas
        import re
        result = re.sub(r',(\s*[}\]])', r'\1', result)
        
        return json.loads(result)
    except Exception as e:
        print(f"AI error: {str(e)}")
        raise

# ---------------------------
# Validation helpers
# ---------------------------
def has_value(val):
    if val is None:
        return False
    if isinstance(val, str):
        val = val.strip().lower()
        return bool(val) and val not in ['null', 'none', 'n/a']
    if isinstance(val, (list, dict)):
        return bool(val)
    return bool(val)

def clean_list(items):
    if not items:
        return []
    return [str(x).strip() for x in items if has_value(x)]

# ---------------------------
# Enhanced Resume Generator with Modern Design
# ---------------------------
class ModernATSResumeGenerator:
    """Production-ready ATS resume generator with modern design"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document()
    
    def setup_document(self):
        """Configure margins and default styles"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
    
    def add_horizontal_line(self):
        """Add a thin horizontal separator line"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run('─' * 100)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(180, 180, 180)
    
    def add_header(self, data):
        """Add professional header with contact info"""
        personal = data.get('personal', {})
        
        if not has_value(personal.get('name')):
            return
        
        # Name - Large and centered
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal['name'].upper())
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.paragraph_format.space_after = Pt(4)
        
        # Contact info line
        contact = []
        if has_value(personal.get('email')):
            contact.append(personal['email'])
        if has_value(personal.get('phone')):
            contact.append(personal['phone'])
        if has_value(personal.get('location')):
            contact.append(personal['location'])
        
        if contact:
            contact_para = self.doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact))
            contact_run.font.size = Pt(10)
            contact_para.paragraph_format.space_after = Pt(3)
        
        # Social links
        links = []
        if has_value(personal.get('linkedin')):
            links.append(f"LinkedIn: {personal['linkedin']}")
        if has_value(personal.get('github')):
            links.append(f"GitHub: {personal['github']}")
        if has_value(personal.get('portfolio')):
            links.append(f"Portfolio: {personal['portfolio']}")
        
        if links:
            links_para = self.doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_run = links_para.add_run(' | '.join(links))
            links_run.font.size = Pt(9)
            links_run.font.color.rgb = RGBColor(0, 102, 204)
            links_para.paragraph_format.space_after = Pt(2)
        
        self.add_horizontal_line()
    
    def add_section_title(self, title):
        """Add centered section header"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(2)
        
        run = para.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add underline
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_para.paragraph_format.space_after = Pt(10)
        line_run = line_para.add_run('─' * 80)
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(100, 100, 100)
    
    def add_summary(self, data):
        """Add professional summary"""
        if not has_value(data.get('summary')):
            return
        
        self.add_section_title('Professional Summary')
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(data['summary'])
        run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(12)
    
    def add_education(self, data):
        """Add education section"""
        education = data.get('education', [])
        if not education:
            return
        
        self.add_section_title('Education')
        
        for edu in education:
            if not has_value(edu.get('institution')):
                continue
            
            # Institution line
            inst_para = self.doc.add_paragraph()
            inst_para.paragraph_format.space_after = Pt(2)
            
            inst_run = inst_para.add_run(edu['institution'])
            inst_run.font.bold = True
            inst_run.font.size = Pt(11)
            
            if has_value(edu.get('location')):
                loc_run = inst_para.add_run(f" — {edu['location']}")
                loc_run.font.size = Pt(10)
                loc_run.font.italic = True
            
            # Degree line
            degree_para = self.doc.add_paragraph()
            degree_para.paragraph_format.space_after = Pt(2)
            
            degree_parts = []
            if has_value(edu.get('degree')):
                degree_parts.append(edu['degree'])
            if has_value(edu.get('field')):
                degree_parts.append(edu['field'])
            
            if degree_parts:
                degree_run = degree_para.add_run(', '.join(degree_parts))
                degree_run.font.size = Pt(10)
                degree_run.font.italic = True
            
            if has_value(edu.get('graduation')):
                grad_run = degree_para.add_run(f" | {edu['graduation']}")
                grad_run.font.size = Pt(10)
            
            # GPA
            if has_value(edu.get('gpa')):
                gpa_para = self.doc.add_paragraph()
                gpa_para.paragraph_format.space_after = Pt(2)
                gpa_run = gpa_para.add_run(f"GPA: {edu['gpa']}")
                gpa_run.font.size = Pt(10)
            
            # Achievements
            achievements = clean_list(edu.get('achievements', []))
            if achievements:
                ach_para = self.doc.add_paragraph()
                ach_para.paragraph_format.space_after = Pt(10)
                ach_run = ach_para.add_run(f"Achievements: {', '.join(achievements)}")
                ach_run.font.size = Pt(10)
            else:
                self.doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    def add_experience(self, data):
        """Add work experience with key achievements"""
        experience = data.get('experience', [])
        if not experience:
            return
        
        self.add_section_title('Professional Experience')
        
        for exp in experience:
            if not has_value(exp.get('company')) or not has_value(exp.get('position')):
                continue
            
            # Company line
            comp_para = self.doc.add_paragraph()
            comp_para.paragraph_format.space_after = Pt(2)
            comp_run = comp_para.add_run(exp['company'])
            comp_run.font.bold = True
            comp_run.font.size = Pt(11)
            
            if has_value(exp.get('location')):
                loc_run = comp_para.add_run(f" — {exp['location']}")
                loc_run.font.size = Pt(10)
                loc_run.font.italic = True
            
            # Position and dates
            pos_para = self.doc.add_paragraph()
            pos_para.paragraph_format.space_after = Pt(4)
            pos_run = pos_para.add_run(exp['position'])
            pos_run.font.size = Pt(10)
            pos_run.font.italic = True
            
            dates = []
            if has_value(exp.get('start')):
                dates.append(exp['start'])
            if has_value(exp.get('end')):
                dates.append(exp['end'])
            
            if dates:
                date_run = pos_para.add_run(f" | {' – '.join(dates)}")
                date_run.font.size = Pt(10)
            
            # Description (if exists)
            if has_value(exp.get('description')):
                desc_para = self.doc.add_paragraph()
                desc_para.paragraph_format.space_after = Pt(4)
                desc_run = desc_para.add_run(exp['description'])
                desc_run.font.size = Pt(10)
                desc_run.font.italic = True
            
            # Key Achievements (most important for ATS!)
            achievements = clean_list(exp.get('achievements', []))
            for achievement in achievements:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.3)
                bullet_para.paragraph_format.space_after = Pt(3)
                bullet_para.paragraph_format.line_spacing = 1.15
                bullet_run = bullet_para.add_run(achievement)
                bullet_run.font.size = Pt(10)
            
            self.doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    def add_projects(self, data):
        """Add projects with achievements"""
        projects = data.get('projects', [])
        if not projects:
            return
        
        self.add_section_title('Projects')
        
        for proj in projects:
            if not has_value(proj.get('name')):
                continue
            
            # Project name
            name_para = self.doc.add_paragraph()
            name_para.paragraph_format.space_after = Pt(2)
            name_run = name_para.add_run(proj['name'])
            name_run.font.bold = True
            name_run.font.size = Pt(11)
            
            # Dates (if exists)
            dates = []
            if has_value(proj.get('start')):
                dates.append(proj['start'])
            if has_value(proj.get('end')):
                dates.append(proj['end'])
            
            if dates:
                date_run = name_para.add_run(f" | {' – '.join(dates)}")
                date_run.font.size = Pt(9)
                date_run.font.italic = True
            
            # Description
            if has_value(proj.get('description')):
                desc_para = self.doc.add_paragraph()
                desc_para.paragraph_format.space_after = Pt(2)
                desc_run = desc_para.add_run(proj['description'])
                desc_run.font.size = Pt(10)
                desc_run.font.italic = True
            
            # Technologies
            technologies = clean_list(proj.get('technologies', []))
            if technologies:
                tech_para = self.doc.add_paragraph()
                tech_para.paragraph_format.space_after = Pt(4)
                tech_run = tech_para.add_run(f"Technologies: {', '.join(technologies)}")
                tech_run.font.size = Pt(9)
                tech_run.font.color.rgb = RGBColor(80, 80, 80)
            
            # Key Achievements
            achievements = clean_list(proj.get('achievements', []))
            for achievement in achievements:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.3)
                bullet_para.paragraph_format.space_after = Pt(3)
                bullet_para.paragraph_format.line_spacing = 1.15
                bullet_run = bullet_para.add_run(achievement)
                bullet_run.font.size = Pt(10)
            
            self.doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    def add_skills(self, data):
        """Add comprehensive skills section"""
        skills = data.get('skills', {})
        if not skills or not any(has_value(v) for v in skills.values()):
            return
        
        self.add_section_title('Skills')
        
        # Technical Skills
        technical = clean_list(skills.get('technical', []))
        if technical:
            self._add_skill_line('Technical Skills', technical)
        
        # Soft Skills
        soft = clean_list(skills.get('soft', []))
        if soft:
            self._add_skill_line('Soft Skills', soft)
        
        # Languages
        languages = clean_list(skills.get('languages', []))
        if languages:
            self._add_skill_line('Languages', languages)
        
        # Tools
        tools = clean_list(skills.get('tools', []))
        if tools:
            self._add_skill_line('Tools & Platforms', tools)
        
        # Additional Skills
        additional = clean_list(skills.get('additional', []))
        if additional:
            self._add_skill_line('Additional Skills', additional)
    
    def _add_skill_line(self, label, items):
        """Helper to add a skill category line"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        label_run = para.add_run(f'{label}: ')
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        items_run = para.add_run(', '.join(items))
        items_run.font.size = Pt(10)
    
    def add_certifications(self, data):
        """Add certifications section"""
        certs = data.get('certifications', [])
        if not certs:
            return
        
        self.add_section_title('Certifications')
        
        for cert in certs:
            if not has_value(cert.get('name')):
                continue
            
            cert_para = self.doc.add_paragraph()
            cert_para.paragraph_format.space_after = Pt(4)
            
            # Certification name
            name_run = cert_para.add_run(cert['name'])
            name_run.font.bold = True
            name_run.font.size = Pt(10)
            
            # Issuer
            if has_value(cert.get('issuer')):
                issuer_run = cert_para.add_run(f" — {cert['issuer']}")
                issuer_run.font.size = Pt(10)
            
            # Date
            if has_value(cert.get('date')):
                date_run = cert_para.add_run(f" | {cert['date']}")
                date_run.font.size = Pt(9)
                date_run.font.italic = True
            
            # Credential ID
            if has_value(cert.get('credential')):
                cred_run = cert_para.add_run(f" (ID: {cert['credential']})")
                cred_run.font.size = Pt(9)
                cred_run.font.color.rgb = RGBColor(100, 100, 100)
    
    def add_leadership(self, data):
        """Add leadership & activities"""
        leadership = data.get('leadership', [])
        if not leadership:
            return
        
        self.add_section_title('Leadership & Activities')
        
        for lead in leadership:
            if not has_value(lead.get('organization')) or not has_value(lead.get('role')):
                continue
            
            # Organization line
            org_para = self.doc.add_paragraph()
            org_para.paragraph_format.space_after = Pt(2)
            org_run = org_para.add_run(lead['organization'])
            org_run.font.bold = True
            org_run.font.size = Pt(11)
            
            if has_value(lead.get('location')):
                loc_run = org_para.add_run(f" — {lead['location']}")
                loc_run.font.size = Pt(10)
                loc_run.font.italic = True
            
            # Role and dates
            role_para = self.doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(4)
            role_run = role_para.add_run(lead['role'])
            role_run.font.size = Pt(10)
            role_run.font.italic = True
            
            dates = []
            if has_value(lead.get('start')):
                dates.append(lead['start'])
            if has_value(lead.get('end')):
                dates.append(lead['end'])
            
            if dates:
                date_run = role_para.add_run(f" | {' – '.join(dates)}")
                date_run.font.size = Pt(10)
            
            # Achievements
            achievements = clean_list(lead.get('achievements', []))
            for achievement in achievements:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.3)
                bullet_para.paragraph_format.space_after = Pt(3)
                bullet_para.paragraph_format.line_spacing = 1.15
                bullet_run = bullet_para.add_run(achievement)
                bullet_run.font.size = Pt(10)
            
            self.doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    def generate(self, data):
        """Generate complete resume"""
        self.add_header(data)
        self.add_summary(data)
        self.add_education(data)
        self.add_experience(data)
        self.add_projects(data)
        self.add_skills(data)
        self.add_certifications(data)
        self.add_leadership(data)
        
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

# ---------------------------
# Routes
# ---------------------------
@app.route('/')
def index():
    return render_template('hybrid_form.html')

@app.route('/upload-and-extract', methods=['POST'])
def upload_and_extract():
    try:
        resume_text = ""
        
        if 'file' in request.files:
            file = request.files['file']
            if file and file.filename and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
                file.save(filepath)
                
                try:
                    if filename.lower().endswith('.pdf'):
                        resume_text = extract_text_from_pdf(filepath)
                    elif filename.lower().endswith('.docx'):
                        resume_text = extract_text_from_docx(filepath)
                    elif filename.lower().endswith('.txt'):
                        with open(filepath, 'r', encoding='utf-8') as f:
                            resume_text = f.read()
                finally:
                    try:
                        os.remove(filepath)
                    except:
                        pass
        
        if not resume_text and request.form.get('text_input'):
            resume_text = request.form.get('text_input')
        
        if not resume_text or len(resume_text.strip()) < 50:
            return jsonify({'error': 'Please provide valid resume content'}), 400
        
        print(f"Processing {len(resume_text)} characters...")
        structured_data = extract_resume_data_with_ai(resume_text)
        print("Extraction successful!")
        return jsonify(structured_data)
    
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    try:
        data = request.get_json(force=True)
        
        personal = data.get('personal', {})
        if not has_value(personal.get('name')):
            return jsonify({'error': 'Name is required'}), 400
        
        print(f"Generating resume for: {personal.get('name')}")
        
        generator = ModernATSResumeGenerator()
        resume_buffer = generator.generate(data)
        
        filename = f"{personal.get('name', 'Resume').replace(' ', '_')}_ATS_Resume.docx"
        
        return send_file(
            resume_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
